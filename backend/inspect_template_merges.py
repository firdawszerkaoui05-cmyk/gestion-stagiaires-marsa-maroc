#!/usr/bin/env python
from docx import Document
from docx.oxml.ns import qn

doc = Document("FICHE ACCUEIL DES STAGIAIRES DE PASSAGE.docx")

def get_cell_merge_info(cell):
    tc = cell._tc
    tcPr = tc.tcPr
    info = {}
    gridSpan = tcPr.xpath('w:gridSpan')
    if gridSpan:
        info['gridSpan'] = int(gridSpan[0].get(qn('w:val')))
    vMerge = tcPr.xpath('w:vMerge')
    if vMerge:
        info['vMerge'] = vMerge[0].get(qn('w:val')) or 'continue'
    return info

print('TABLES', len(doc.tables))
for t_idx, table in enumerate(doc.tables):
    print(f'\nTABLE {t_idx}')
    for r_idx, row in enumerate(table.rows):
        if r_idx in [2,3,4,5,6,7,9,10,11,12,14,15,16,17]:
            print(f'ROW {r_idx}')
            for c_idx, cell in enumerate(row.cells):
                text = cell.text.replace('\n','\\n')
                info = get_cell_merge_info(cell)
                if text.strip() or info:
                    print(f'  CELL {c_idx}: text={repr(text)} info={info}')
