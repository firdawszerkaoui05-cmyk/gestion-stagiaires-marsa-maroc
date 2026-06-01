#!/usr/bin/env python
# -*- coding: utf-8 -*-
from docx import Document

doc = Document("FICHE ACCUEIL DES STAGIAIRES DE PASSAGE.docx")
print('Tables:', len(doc.tables))
for t_idx, table in enumerate(doc.tables):
    print('\nTABLE', t_idx)
    for r_idx, row in enumerate(table.rows):
        row_text = [cell.text.replace('\n','\\n') for cell in row.cells]
        if any(cell.strip() for cell in row_text):
            print(f'ROW {r_idx}:')
            for c_idx, text in enumerate(row_text):
                if text.strip():
                    print(f'  CELL {c_idx}: {repr(text)}')
