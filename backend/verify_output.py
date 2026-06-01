#!/usr/bin/env python
from docx import Document

doc = Document('uploads/TEST_FICHE_OFFICIELLE.docx')
print('Tables:', len(doc.tables))
for t_idx, table in enumerate(doc.tables):
    print('\nTABLE', t_idx)
    for r_idx, row in enumerate(table.rows):
        if r_idx in [2,4,6,7,9,11,12,14,15,16,17]:
            texts = [cell.text.replace('\n','\\n') for cell in row.cells]
            print(f'ROW {r_idx}:', texts[:5])
