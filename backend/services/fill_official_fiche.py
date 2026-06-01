#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Service de remplissage automatique de la fiche officielle Marsa Maroc.
Utilise la fiche officielle "FICHE ACCUEIL DES STAGIAIRES DE PASSAGE".
"""

from datetime import datetime
import os
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


def format_date(date_str):
    """Convertit une date au format DD/MM/YYYY."""
    if not date_str:
        return ""

    try:
        if isinstance(date_str, str) and len(date_str) == 10 and date_str.count("-") == 2:
            date_obj = datetime.strptime(date_str, "%Y-%m-%d")
            return date_obj.strftime("%d/%m/%Y")
        return date_str
    except (ValueError, AttributeError):
        return date_str if date_str else ""


def clean_cell_value(value):
    """Normalise les valeurs injectees sans ajouter de retours a la ligne."""
    if value is None:
        return ""

    value = str(value).replace("\r", " ").replace("\n", " ").replace("\t", " ")
    return re.sub(r"\s+", " ", value).strip()


def font_size_for_value(value):
    """Reduit progressivement la police pour garder le tableau stable."""
    length = len(value)

    if length <= 18:
        return 11
    if length <= 32:
        return 10
    if length <= 48:
        return 9
    if length <= 68:
        return 8
    if length <= 95:
        return 7
    return 6


def enable_single_line_cell(cell):
    """
    Demande a Word de ne pas enrouler le texte et de l'ajuster dans la cellule.
    Ces proprietes conservent la grille originale du modele officiel.
    """
    tc_pr = cell._tc.get_or_add_tcPr()

    for tag in ("w:noWrap", "w:tcFitText"):
        if tc_pr.find(qn(tag)) is None:
            tc_pr.append(OxmlElement(tag))


def style_compact_paragraph(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1


def fill_official_fiche(stagiaire, affectation, output_path):
    """
    Remplit la fiche officielle Marsa Maroc avec les donnees du stagiaire.

    Args:
        stagiaire: Dict avec (nom, prenom, telephone, etablissement, profil,
            type_stage, date_debut, date_fin, recommande)
        affectation: Dict avec (division, encadrant, theme_stage)
        output_path: Chemin de sortie du fichier rempli

    Returns:
        output_path si succes, None sinon
    """

    try:
        template_path = "FICHE ACCUEIL DES STAGIAIRES DE PASSAGE.docx"

        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Template non trouve: {template_path}")

        doc = Document(template_path)

        if len(doc.tables) < 2:
            raise ValueError("Le template n'a pas la structure attendue (min 2 tableaux)")

        main_table = doc.tables[1]

        nom = clean_cell_value(stagiaire.get("nom", "")).upper()
        prenom = clean_cell_value(stagiaire.get("prenom", "")).upper()
        telephone = clean_cell_value(stagiaire.get("telephone", ""))
        etablissement = clean_cell_value(stagiaire.get("etablissement", ""))
        profil = clean_cell_value(stagiaire.get("profil", ""))
        type_stage = clean_cell_value(stagiaire.get("type_stage", "Passage"))
        recommande = clean_cell_value(stagiaire.get("recommande", ""))
        date_debut = format_date(stagiaire.get("date_debut", ""))
        date_fin = format_date(stagiaire.get("date_fin", ""))

        division = clean_cell_value(affectation.get("division", "")) if affectation else ""
        encadrant = clean_cell_value(affectation.get("encadrant", "")) if affectation else ""
        theme_stage = clean_cell_value(affectation.get("theme_stage", "")) if affectation else ""

        def append_value_to_label(row_idx, value, paragraph_idx=None):
            value = clean_cell_value(value)
            if not value:
                return

            cell = main_table.rows[row_idx].cells[0]
            enable_single_line_cell(cell)

            paragraphs = cell.paragraphs
            if paragraph_idx is None:
                paragraph_idx = next(
                    (idx for idx in range(len(paragraphs) - 1, -1, -1) if paragraphs[idx].text.strip()),
                    0,
                )

            paragraph = paragraphs[paragraph_idx]
            style_compact_paragraph(paragraph)

            separator = " " if paragraph.text and not paragraph.text.endswith((" ", "\u00a0", "\t")) else ""
            run = paragraph.add_run(f"{separator}{value}")
            run.font.size = Pt(font_size_for_value(value))
            run.font.name = "Arial"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")

        def fill_stage_period(row_idx, start_date, end_date):
            if not start_date and not end_date:
                return

            append_value_to_label(row_idx, f"Du {start_date} au {end_date}")

        append_value_to_label(2, nom)
        append_value_to_label(4, prenom)
        append_value_to_label(6, telephone)
        append_value_to_label(7, etablissement, paragraph_idx=1)
        append_value_to_label(9, profil)
        append_value_to_label(11, type_stage)
        append_value_to_label(12, recommande, paragraph_idx=2)
        append_value_to_label(14, division, paragraph_idx=3)
        append_value_to_label(15, encadrant)
        append_value_to_label(16, theme_stage)
        fill_stage_period(17, date_debut, date_fin)

        output_dir = os.path.dirname(output_path)
        if output_dir:
            os.makedirs(output_dir, exist_ok=True)

        doc.save(output_path)

        print(f"Fiche remplie avec succes : {output_path}")
        return output_path

    except Exception as e:
        print(f"Erreur lors du remplissage : {str(e)}")
        return None


if __name__ == "__main__":
    stagiaire = {
        "nom": "Zerkaoui",
        "prenom": "Firdaws",
        "telephone": "0674182136",
        "etablissement": "FST Settat",
        "profil": "Licence en systeme d'information",
        "type_stage": "Projet de fin d'etude",
        "date_debut": "2026-04-01",
        "date_fin": "2026-06-01",
        "recommande": "Par site web",
    }

    affectation = {
        "division": "Informatique",
        "encadrant": "Mme Amina Belkadi",
        "theme_stage": "Developpement Backend",
    }

    output = "uploads/TEST_FICHE_OFFICIELLE.docx"
    fill_official_fiche(stagiaire, affectation, output)
