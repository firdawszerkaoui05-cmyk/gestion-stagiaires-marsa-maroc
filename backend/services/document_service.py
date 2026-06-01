# -*- coding: utf-8 -*-
"""
Service de génération des documents Word - FICHE D'ACCUEIL DES STAGIAIRES
Utilise python-docx pour remplir automatiquement les données du stagiaire et l'affectation IA
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os


def format_date(date_str):
    """
    Convertit une date au format DD/MM/YYYY
    
    Args:
        date_str: Date en format ISO (YYYY-MM-DD) ou déjà formatée
        
    Returns:
        Date formatée en DD/MM/YYYY ou date_str si impossible
    """
    if not date_str:
        return "Non spécifiée"
    
    try:
        if isinstance(date_str, str):
            # Essayer le format ISO
            if len(date_str) == 10 and date_str.count('-') == 2:
                date_obj = datetime.strptime(date_str, "%Y-%m-%d")
                return date_obj.strftime("%d/%m/%Y")
            # Retourner tel quel si déjà formaté
            return date_str
    except (ValueError, AttributeError):
        return date_str if date_str else "Non spécifiée"


def add_section_title(doc, title):
    """Ajoute un titre de section formaté"""
    heading = doc.add_paragraph()
    heading.style = 'Heading 1'
    run = heading.add_run(title)
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)  # Bleu foncé
    
    
def add_info_row(table, label, value):
    """
    Ajoute une ligne d'information dans un tableau
    
    Args:
        table: Tableau Word
        label: Libellé (colonne 1)
        value: Valeur (colonne 2)
    """
    row = table.add_row()
    row.cells[0].text = label
    row.cells[1].text = str(value) if value else "Non spécifiée"


def generate_fiche_accueil(stagiaire, affectation=None):
    """
    Génère la FICHE D'ACCUEIL DES STAGIAIRES en Word avec:
    - Données du stagiaire
    - Affectation IA (si disponible)
    - Formatage professionnel Marsa Maroc
    
    Args:
        stagiaire: Dict avec clés (nom, prenom, email, telephone, etablissement, 
                                   profil, type_stage, date_debut, date_fin)
        affectation: Dict avec clés (division, encadrant, theme_stage, domaine, 
                                     justification) [optionnel]
                                     
    Returns:
        Document Word (Document object from python-docx)
    """
    doc = Document()
    
    # ================= ENTÊTE =================
    # Titre principal
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("FICHE D'ACCUEIL DES STAGIAIRES")
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 102, 51)  # Vert Marsa Maroc
    
    # Sous-titre
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Marsa Maroc")
    run.font.size = Pt(11)
    run.font.italic = True
    
    # Ligne de séparation
    doc.add_paragraph("_" * 80)
    
    # ================= SECTION STAGIAIRE =================
    add_section_title(doc, "INFORMATIONS DU STAGIAIRE")
    
    # Tableau pour les infos stagiaire
    table_stagiaire = doc.add_table(rows=1, cols=2)
    table_stagiaire.style = 'Light Grid Accent 1'
    
    # En-tête du tableau
    header_cells = table_stagiaire.rows[0].cells
    header_cells[0].text = "Champ"
    header_cells[1].text = "Valeur"
    
    # Remplir les données stagiaire
    add_info_row(table_stagiaire, "Nom", stagiaire.get("nom", ""))
    add_info_row(table_stagiaire, "Prénom", stagiaire.get("prenom", ""))
    add_info_row(table_stagiaire, "Email", stagiaire.get("email", ""))
    add_info_row(table_stagiaire, "Téléphone", stagiaire.get("telephone", ""))
    add_info_row(table_stagiaire, "Établissement", stagiaire.get("etablissement", ""))
    add_info_row(table_stagiaire, "Profil / Niveau", stagiaire.get("profil", ""))
    add_info_row(table_stagiaire, "Type de stage", stagiaire.get("type_stage", ""))
    
    # Bloc Périodes
    doc.add_paragraph()
    periods = doc.add_paragraph()
    periods.add_run("Période de stage : ").bold = True
    date_debut = format_date(stagiaire.get("date_debut", ""))
    date_fin = format_date(stagiaire.get("date_fin", ""))
    periods.add_run(f"Du {date_debut} Au {date_fin}")
    
    # ================= SECTION AFFECTATION =================
    if affectation:
        doc.add_paragraph()
        add_section_title(doc, "AFFECTATION À L'ENTITÉ D'ACCUEIL")
        
        # Tableau pour l'affectation
        table_affectation = doc.add_table(rows=1, cols=2)
        table_affectation.style = 'Light Grid Accent 1'
        
        # En-tête du tableau
        header_cells = table_affectation.rows[0].cells
        header_cells[0].text = "Élément"
        header_cells[1].text = "Détail"
        
        # Remplir les données d'affectation
        add_info_row(table_affectation, "Division", affectation.get("division", ""))
        add_info_row(table_affectation, "Encadrant", affectation.get("encadrant", ""))
        add_info_row(table_affectation, "Thème de stage", affectation.get("theme_stage", ""))
        add_info_row(table_affectation, "Domaine", affectation.get("domaine", ""))
        
        # Justification
        if affectation.get("justification"):
            doc.add_paragraph()
            just = doc.add_paragraph()
            just.add_run("Justification de l'affectation : ").bold = True
            just.add_run(affectation.get("justification"))
    
    # ================= SECTION NOTES =================
    doc.add_paragraph()
    doc.add_paragraph()
    notes = doc.add_paragraph()
    notes.add_run("Remarques et notes : ").bold = True
    doc.add_paragraph("_" * 80)
    doc.add_paragraph()
    doc.add_paragraph()
    
    # ================= PIED DE PAGE =================
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(f"Document généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')}")
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(128, 128, 128)  # Gris
    
    return doc


def save_fiche_accueil(stagiaire, affectation, filepath):
    """
    Génère et sauvegarde la fiche d'accueil Word
    
    Args:
        stagiaire: Dict avec les données du stagiaire
        affectation: Dict avec les données d'affectation (optionnel)
        filepath: Chemin où sauvegarder le fichier
    """
    doc = generate_fiche_accueil(stagiaire, affectation)
    
    # Créer le répertoire s'il n'existe pas
    os.makedirs(os.path.dirname(filepath), exist_ok=True)
    
    # Sauvegarder
    doc.save(filepath)
    return filepath
