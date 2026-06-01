import os
from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from models.stagiaire import db, Stagiaire
from services.document_scan import scan_documents

UPLOAD_FOLDER = "uploads"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


def generate_word(stagiaire, filepath):
    doc = Document()
    doc.add_heading("FICHE D’ACCUEIL STAGIAIRE", 0)
    doc.add_paragraph(f"Nom : {stagiaire['nom']}")
    doc.add_paragraph(f"Prénom : {stagiaire['prenom']}")
    doc.add_paragraph(f"Téléphone : {stagiaire['telephone']}")
    doc.add_paragraph(f"Etablissement : {stagiaire['etablissement']}")
    doc.add_paragraph(f"Profil : {stagiaire['profil']}")
    doc.add_paragraph(f"Type de stage : {stagiaire['type_stage']}")
    doc.add_paragraph(f"Recommandé par : {stagiaire['recommande']}")
    doc.save(filepath)


def generate_pdf(stagiaire, filepath):
    doc = SimpleDocTemplate(filepath)
    styles = getSampleStyleSheet()
    content = [
        Paragraph("FICHE STAGIAIRE", styles["Title"]),
        Paragraph(f"Nom: {stagiaire['nom']}", styles["Normal"]),
        Paragraph(f"Prénom: {stagiaire['prenom']}", styles["Normal"]),
        Paragraph(f"Téléphone: {stagiaire['telephone']}", styles["Normal"]),
        Paragraph(f"Etablissement: {stagiaire['etablissement']}", styles["Normal"]),
        Paragraph(f"Profil: {stagiaire['profil']}", styles["Normal"]),
        Paragraph(f"Type de stage: {stagiaire['type_stage']}", styles["Normal"]),
        Paragraph(f"Recommandé par: {stagiaire['recommande']}", styles["Normal"]),
    ]
    doc.build(content)


def save_files(file_list):
    documents = ["convention", "assurance", "cin", "demande_stage"]
    saved_files = {}
    for idx, f in enumerate(file_list):
        if not f:
            continue
        filename = f.filename
        path = os.path.join(UPLOAD_FOLDER, filename)
        f.save(path)
        key = documents[idx] if idx < len(documents) else f"fichier_{idx}"
        saved_files[key] = filename
    return saved_files


def add_stagiaire(data, file_list):
    saved_files = save_files(file_list)

    # Créer un nouveau stagiaire en base de données
    stagiaire = Stagiaire(
        nom=data.get("nom"),
        prenom=data.get("prenom"),
        email=data.get("email"),
        telephone=data.get("telephone"),
        etablissement=data.get("etablissement"),
        profil=data.get("profil_niveau"),
        type_stage=data.get("type_stage"),
        recommande=data.get("recommande_par"),
        decision="En attente",
        pdf=f"fiche_temp.pdf",
        word=f"fiche_temp.docx"
    )

    db.session.add(stagiaire)
    db.session.commit()

    # Mise à jour des noms de fichiers avec l'ID généré
    stagiaire.pdf = f"fiche_{stagiaire.id}.pdf"
    stagiaire.word = f"fiche_{stagiaire.id}.docx"

    # Générer les fichiers
    stagiaire_dict = stagiaire.to_dict()
    generate_pdf(stagiaire_dict, os.path.join(UPLOAD_FOLDER, stagiaire.pdf))
    generate_word(stagiaire_dict, os.path.join(UPLOAD_FOLDER, stagiaire.word))

    db.session.commit()

    scan_report = scan_documents(saved_files)

    return {
        "stagiaire": stagiaire.to_dict(),
        "scan_report": scan_report,
    }


def get_stagiaires():
    stagiaires = Stagiaire.query.all()
    return [s.to_dict() for s in stagiaires]


def set_decision(stagiaire_id, decision):
    stagiaire = Stagiaire.query.get(stagiaire_id)
    if stagiaire:
        stagiaire.decision = decision
        db.session.commit()
        return stagiaire.to_dict()
    return None


def get_stagiaire_by_id(stagiaire_id):
    stagiaire = Stagiaire.query.get(stagiaire_id)
    if stagiaire:
        return stagiaire.to_dict()
    return None


def analyse_candidature(data):
    return {
        "valid": True,
        "message": "Analyse IA simulée : dossier cohérent.",
        "details": {
            "cv": True,
            "cin": True,
            "lettre_motivation": True
        }
    }
