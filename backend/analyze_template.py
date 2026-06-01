#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Analyse de la fiche officielle Marsa Maroc"""

from docx import Document

# Charger le modèle officiel
doc = Document("FICHE ACCUEIL DES STAGIAIRES DE PASSAGE.docx")

print("=" * 80)
print("ANALYSE DU MODÈLE - FICHE D'ACCUEIL DES STAGIAIRES DE PASSAGE")
print("=" * 80)

print("\n[STRUCTURE]")
print(f"Nombre de paragraphes: {len(doc.paragraphs)}")
print(f"Nombre de tableaux: {len(doc.tables)}")

print("\n[CONTENU DES PARAGRAPHES (premiers 40)]")
for i, para in enumerate(doc.paragraphs[:40]):
    if para.text.strip():
        print(f"{i}: {para.text[:100]}")

print("\n[STRUCTURE DES TABLEAUX]")
for t_idx, table in enumerate(doc.tables):
    print(f"\nTableau {t_idx}:")
    print(f"  Lignes: {len(table.rows)}")
    print(f"  Colonnes: {len(table.columns)}")
    print(f"  Contenu des lignes:")
    for r_idx, row in enumerate(table.rows):
        cells_text = [cell.text[:40].strip() for cell in row.cells]
        print(f"    Ligne {r_idx}: {cells_text}")

print("\n" + "=" * 80)
print("FIN ANALYSE")
print("=" * 80)
